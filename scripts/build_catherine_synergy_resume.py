from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/resume/Catherine_Zharoff_Synergy_Claims_Resume.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(24, 55, 82)
GRAY = RGBColor(82, 91, 99)
BLACK = RGBColor(25, 25, 25)


def font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_bottom_border(paragraph, color="183752", size="10", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    font(r, size=10.5, bold=True, color=NAVY)
    add_bottom_border(p, size="6", space="3")
    return p


def add_bullet(doc, text, after=2.5):
    p = doc.add_paragraph(style="Resume Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    font(r, size=9.8)
    return p


def add_job_header(doc, title, dates, company, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    keep_with_next(p)
    r = p.add_run(title)
    font(r, size=10.5, bold=True)
    r = p.add_run(f"  |  {dates}")
    font(r, size=10, bold=True, color=GRAY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    keep_with_next(p2)
    r = p2.add_run(company)
    font(r, size=10, bold=True, color=NAVY)
    r = p2.add_run(f"  |  {location}")
    font(r, size=10, italic=True, color=GRAY)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.55)
sec.bottom_margin = Inches(0.45)
sec.left_margin = Inches(0.72)
sec.right_margin = Inches(0.72)
sec.header_distance = Inches(0.3)
sec.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.05

bullet = styles.add_style("Resume Bullet", 1)
bullet.font.name = "Arial"
bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
bullet.font.size = Pt(9.8)
bullet.paragraph_format.left_indent = Inches(0.23)
bullet.paragraph_format.first_line_indent = Inches(-0.14)
bullet.paragraph_format.space_before = Pt(0)
bullet.paragraph_format.space_after = Pt(2.5)
bullet.paragraph_format.line_spacing = 1.02

# Real bullet numbering definition.
numbering = doc.part.numbering_part.element
abstract = OxmlElement("w:abstractNum")
abstract.set(qn("w:abstractNumId"), "42")
multi = OxmlElement("w:multiLevelType")
multi.set(qn("w:val"), "singleLevel")
abstract.append(multi)
lvl = OxmlElement("w:lvl")
lvl.set(qn("w:ilvl"), "0")
start = OxmlElement("w:start")
start.set(qn("w:val"), "1")
lvl.append(start)
num_fmt = OxmlElement("w:numFmt")
num_fmt.set(qn("w:val"), "bullet")
lvl.append(num_fmt)
lvl_text = OxmlElement("w:lvlText")
lvl_text.set(qn("w:val"), "•")
lvl.append(lvl_text)
lvl_jc = OxmlElement("w:lvlJc")
lvl_jc.set(qn("w:val"), "left")
lvl.append(lvl_jc)
ppr = OxmlElement("w:pPr")
tabs = OxmlElement("w:tabs")
tab = OxmlElement("w:tab")
tab.set(qn("w:val"), "num")
tab.set(qn("w:pos"), "331")
tabs.append(tab)
ppr.append(tabs)
ind = OxmlElement("w:ind")
ind.set(qn("w:left"), "331")
ind.set(qn("w:hanging"), "202")
ppr.append(ind)
lvl.append(ppr)
abstract.append(lvl)
numbering.append(abstract)
num = OxmlElement("w:num")
num.set(qn("w:numId"), "42")
abstract_id = OxmlElement("w:abstractNumId")
abstract_id.set(qn("w:val"), "42")
num.append(abstract_id)
numbering.append(num)

for p in []:
    pass

# Header block (customer-pack inspired, ATS-safe single column).
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
r = p.add_run("CATHERINE ZHAROFF")
font(r, size=20, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
r = p.add_run("VETERINARY CLAIMS & MEDICAL RECORDS CANDIDATE")
font(r, size=10.5, bold=True, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Sumner, WA  |  (425) 802-3316  |  catherine.zharoff@gmail.com")
font(r, size=9.5)
add_bottom_border(p, size="8", space="5")

section_heading(doc, "Professional Summary")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run(
    "Veterinary professional with 20+ years of small-animal clinical experience and a strong foundation in "
    "medical documentation, disease processes, pharmaceuticals, diagnostics, treatment workflows, and client "
    "education. Experienced reviewing patient histories, recording clinical information, coordinating prescriptions, "
    "and working accurately across high-volume cases. Proficient in Cornerstone and AVImark; prepared to apply clinical "
    "expertise to remote pet-insurance claims review and file documentation."
)
font(r, size=10)

section_heading(doc, "Claims-Relevant Qualifications")
qualifications = [
    "Veterinary medical records, chart notation, patient histories, invoices, and treatment documentation",
    "Disease processes, acute and chronic conditions, medications, vaccines, diagnostics, and laboratory results",
    "Cornerstone and AVImark practice-management software; comfortable learning new technology and workflows",
    "Clear written and verbal communication with veterinarians, clinic teams, and pet owners",
    "High-volume, multi-line telephone experience; client education regarding medications, treatment, and aftercare",
    "Detail-oriented case follow-through, prioritization, confidentiality, and independent judgment",
]
for text in qualifications:
    p = add_bullet(doc, text, after=2)
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42")
    num_pr.append(ilvl)
    num_pr.append(num_id)

section_heading(doc, "Professional Experience")
add_job_header(doc, "Veterinary Assistant", "2019–2021; April 2023–Present", "Timber Ridge Animal Hospital", "Bonney Lake, WA")
timber = [
    "Support veterinarians during examinations, surgeries, and dental procedures while maintaining accurate, timely patient documentation.",
    "Review patient histories, clinical findings, diagnostics, medications, and treatment instructions to support continuity of care.",
    "Prepare and administer vaccines and medications under veterinarian direction; apply extensive knowledge of pharmaceuticals, dosing workflows, and patient safety.",
    "Perform diagnostic radiography and assist with laboratory and procedural workflows, recognizing information relevant to acute and chronic conditions.",
    "Educate clients about medications, treatment plans, and aftercare; access and update patient records, treatments, and communications.",
]
for text in timber:
    p = add_bullet(doc, text)
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), "42")
    num_pr.append(ilvl); num_pr.append(num_id)

add_job_header(doc, "Assistant Manager", "2021–2022 (9 months)", "SJC Property Management Group", "Washington")
sjc = [
    "Handled client communications and administrative work with accuracy, professionalism, and careful follow-through.",
    "Coordinated and processed utility payments involving multiple cities and counties, maintaining organized records and meeting deadlines.",
]
for text in sjc:
    p = add_bullet(doc, text)
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), "42")
    num_pr.append(ilvl); num_pr.append(num_id)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Family Caregiving  |  2022–April 2023")
font(r, size=9.5, bold=True, color=GRAY)

add_job_header(doc, "Veterinary Assistant / Staff Lead", "January 2001–2018", "Animal Hospital of Newport Hills", "Newcastle, WA")
newport = [
    "Supported seven veterinarians in a busy walk-in small-animal practice, managing competing case priorities and documenting clinical information accurately.",
    "Recorded chart notes, coordinated case details, reviewed procedure invoices, and maintained clear information across patient and client workflows.",
    "Collected blood, urine, fecal, tissue, and other specimens; performed in-house urinalysis and used IDEXX Catalyst and ProCyte diagnostic equipment.",
    "Assisted with surgical induction, preparation, anesthesia monitoring, recovery, dental care, radiography, emergency triage, and routine examinations.",
    "Filled prescriptions, coordinated refills, monitored inventory, maintained patient information, and supported multi-line telephone communication and follow-up.",
]
for text in newport:
    p = add_bullet(doc, text)
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), "42")
    num_pr.append(ilvl); num_pr.append(num_id)

section_heading(doc, "Education")
education = [("Veterinary Assistant Program", "Bellingham Technical College — Bellingham, WA")]
for credential, school in education:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{credential}: ")
    font(r, size=9.5, bold=True)
    r = p.add_run(school)
    font(r, size=9.5)

# Apply numPr to every bullet style paragraph in case style inheritance changes.
for p in doc.paragraphs:
    if p.style.name == "Resume Bullet" and p._p.pPr.numPr is None:
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), "42")
        num_pr.append(ilvl); num_pr.append(num_id)

# Quiet footer.
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
r = fp.add_run("Catherine Zharoff  |  Veterinary Claims Candidate")
font(r, size=8, color=GRAY)

doc.core_properties.title = "Catherine Zharoff - Synergy Pet Group Claims Adjuster Resume"
doc.core_properties.subject = "Targeted resume for Claims Adjuster"
doc.core_properties.author = "Catherine Zharoff"
doc.core_properties.keywords = "veterinary claims, medical records, Cornerstone, AVImark, pet insurance"

doc.save(OUT)
print(OUT.resolve())
